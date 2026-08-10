#!/usr/bin/env python3
"""
orcamento_masterleds.py — Preenche o modelo de orçamento da MasterLeds.

Usa templates/orcamento-masterleds.docx como base (que já traz logo no
cabeçalho, endereço no rodapé e logos de parceiros) e preenche apenas o
texto do corpo: cliente, datas, local, itens de equipamento, valor total e
forma de pagamento. As imagens e a identidade visual são preservadas porque
ficam no cabeçalho/rodapé, que não são tocados.

Uso:
    python orcamento_masterleds.py dados.json --out "orcamento-cliente"

Ou passando os campos direto (dados.json tem prioridade se ambos existirem):
    python orcamento_masterleds.py --cliente "Fulano" --valor-total "R$ 10.000,00" \\
        --data-evento "20/09/2026" --item "01 Painel de Led ..." --item "..."

Formato do dados.json (todos os campos são opcionais):
{
  "cliente": "Empresa X",
  "data_evento": "20/09/2026",
  "data_montagem": "19/09/2026",
  "local": "Centro de Eventos - Porto Alegre/RS",
  "itens": [
    "01 Painel de Led PH2.9mm Indoor Absen - Tamanho 4x2m",
    "01 Processador VX1000",
    "Transportes - Equipamentos e Equipe"
  ],
  "valor_total": "R$ 12.500,00",
  "forma_pagamento": "30/60 DD Boleto bancário"
}

O resultado é um .docx com a marca MasterLeds. Para gerar o PDF, abra no Word
e use "Salvar como / Exportar para PDF" (o layout com logo é preservado).

Dependência:
    pip install python-docx
"""
import argparse
import json
import os
import sys


TEMPLATE_REL = os.path.join("..", "templates", "orcamento-masterleds.docx")


def _default_template():
    here = os.path.dirname(os.path.abspath(__file__))
    return os.path.normpath(os.path.join(here, TEMPLATE_REL))


def _set_text(paragraph, text):
    """Substitui o texto do parágrafo mantendo a formatação do 1º run."""
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(text)


def _delete(paragraph):
    el = paragraph._element
    el.getparent().remove(el)


def _find(paragraphs, predicate):
    for idx, p in enumerate(paragraphs):
        if predicate(p.text.strip()):
            return idx
    return None


def fill_orcamento(data, template_path, out_path):
    try:
        from docx import Document
    except ImportError:
        print("ERRO: falta 'python-docx'. Rode: pip install python-docx", file=sys.stderr)
        raise

    doc = Document(template_path)
    paras = doc.paragraphs

    # ---- Campos de cabeçalho (label + valor na mesma linha) ----
    def fill_label(prefix, value):
        if value is None:
            return
        idx = _find(paras, lambda t: t.lower().startswith(prefix.lower()))
        if idx is not None:
            # Mantém o rótulo original (ex.: "Cliente: ") e acrescenta o valor.
            label = paras[idx].text
            if not label.rstrip().endswith(":"):
                # ex.: "Cliente: " -> preserva o que vem antes do 1º ":"
                base = label.split(":", 1)[0] + ": "
            else:
                base = label.rstrip() + " "
            _set_text(paras[idx], base + str(value))

    fill_label("cliente", data.get("cliente"))
    fill_label("data do evento", data.get("data_evento"))
    fill_label("data de montagem", data.get("data_montagem"))
    fill_label("local", data.get("local"))

    # ---- Valor total ----
    if data.get("valor_total") is not None:
        idx = _find(paras, lambda t: t.lower().startswith("valor") and "total" in t.lower())
        if idx is not None:
            _set_text(paras[idx], "Valor Total: " + str(data["valor_total"]))

    # ---- Forma de pagamento (opcional; mantém a padrão se não informado) ----
    if data.get("forma_pagamento"):
        idx = _find(paras, lambda t: t.lower().startswith("forma de pagamento"))
        if idx is not None:
            _set_text(paras[idx], "Forma de pagamento: " + str(data["forma_pagamento"]))

    # ---- Itens de equipamento ----
    # Substitui o bloco de itens (parágrafos não vazios entre "Local:" e
    # "Valor Total") pela lista fornecida, preservando o estilo Normal.
    itens = data.get("itens")
    if itens:
        paras = doc.paragraphs  # re-lê após edições acima
        local_idx = _find(paras, lambda t: t.lower().startswith("local"))
        valor_idx = _find(paras, lambda t: t.lower().startswith("valor") and "total" in t.lower())
        if local_idx is not None and valor_idx is not None and valor_idx > local_idx:
            # Remove itens antigos (parágrafos Normal não vazios no intervalo).
            for p in paras[local_idx + 1:valor_idx]:
                if p.text.strip() and p.style.name == "Normal":
                    _delete(p)
            # Insere os novos itens antes do parágrafo "Valor Total".
            ref = doc.paragraphs[_find(doc.paragraphs,
                                       lambda t: t.lower().startswith("valor") and "total" in t.lower())]
            for item in itens:
                new_p = ref.insert_paragraph_before(str(item), style="Normal")
                _ = new_p

    doc.save(out_path)
    return out_path


def _load_data(args):
    data = {}
    if args.json:
        with open(args.json, "r", encoding="utf-8") as f:
            data = json.load(f)
    # Campos via linha de comando complementam/sobrescrevem o JSON.
    for key, val in (("cliente", args.cliente), ("data_evento", args.data_evento),
                     ("data_montagem", args.data_montagem), ("local", args.local),
                     ("valor_total", args.valor_total),
                     ("forma_pagamento", args.forma_pagamento)):
        if val is not None:
            data[key] = val
    if args.item:
        data["itens"] = args.item
    return data


def main():
    ap = argparse.ArgumentParser(description="Preenche o orçamento modelo da MasterLeds.")
    ap.add_argument("json", nargs="?", help="Arquivo JSON com os dados (opcional)")
    ap.add_argument("--out", help="Nome base de saída (sem extensão)", default="orcamento-masterleds")
    ap.add_argument("--outdir", default=".", help="Pasta de saída")
    ap.add_argument("--template", default=None, help="Caminho do modelo .docx")
    ap.add_argument("--cliente")
    ap.add_argument("--data-evento", dest="data_evento")
    ap.add_argument("--data-montagem", dest="data_montagem")
    ap.add_argument("--local")
    ap.add_argument("--valor-total", dest="valor_total")
    ap.add_argument("--forma-pagamento", dest="forma_pagamento")
    ap.add_argument("--item", action="append", help="Item de equipamento (repita a flag)")
    args = ap.parse_args()

    template = args.template or _default_template()
    if not os.path.exists(template):
        print(f"ERRO: modelo não encontrado: {template}", file=sys.stderr)
        sys.exit(1)

    data = _load_data(args)
    os.makedirs(args.outdir, exist_ok=True)
    out_path = os.path.join(args.outdir, args.out + ".docx")
    fill_orcamento(data, template, out_path)
    print("Gerado:", os.path.abspath(out_path))


if __name__ == "__main__":
    main()
