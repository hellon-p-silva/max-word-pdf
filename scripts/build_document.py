#!/usr/bin/env python3
"""
build_document.py — Gera .docx e/ou .pdf a partir de um arquivo Markdown.

Uso:
    python build_document.py CONTEUDO.md --out "nome-do-arquivo" --docx --pdf

Flags:
    --docx        Gera arquivo .docx (Word)
    --pdf         Gera arquivo .pdf
    --out NOME    Nome base dos arquivos de saída (sem extensão).
                  Se omitido, usa o nome do .md de entrada.
    --outdir DIR  Pasta de saída (padrão: pasta atual).

Se nem --docx nem --pdf forem passados, gera AMBOS.

Markdown suportado (mantenha simples):
    # Título              -> título do documento (centralizado)
    ## Seção              -> cabeçalho de seção
    ### Subseção          -> cabeçalho menor
    - item / * item       -> lista com marcadores
    1. item               -> lista numerada
    parágrafo em texto    -> parágrafo normal (linhas em branco separam)
    **negrito**           -> negrito inline
    ---                   -> quebra de página

    Tabelas (ótimas para orçamentos):
        | Item        | Qtd | Valor    |
        | ----------- | --: | -------: |
        | Consultoria |   1 | R$ 500   |
        | Suporte     |   2 | R$ 200   |
    A 2ª linha (traços) define as colunas; use ":" para alinhar
    (`:--` esquerda, `--:` direita, `:--:` centro). O cabeçalho ganha
    destaque e as linhas ficam zebradas.

Dependências:
    pip install python-docx reportlab
"""
import argparse
import os
import re
import sys


# ----------------------------- Parser de Markdown -----------------------------

def _split_row(line):
    """Divide uma linha de tabela em células, ignorando os | das bordas."""
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def _is_separator_row(line):
    """Detecta a linha separadora de tabela: | --- | :--: | --- |."""
    if "-" not in line:
        return False
    cells = _split_row(line)
    if not cells:
        return False
    return all(re.fullmatch(r":?-+:?", c.strip()) for c in cells if c.strip() != "") \
        and any(c.strip() for c in cells)


def _parse_align(sep_line):
    """Lê o alinhamento de cada coluna a partir da linha separadora."""
    aligns = []
    for c in _split_row(sep_line):
        c = c.strip()
        left, right = c.startswith(":"), c.endswith(":")
        if left and right:
            aligns.append("center")
        elif right:
            aligns.append("right")
        elif left:
            aligns.append("left")
        else:
            aligns.append("")
    return aligns


def parse_markdown(text):
    """Converte markdown simples numa lista de blocos estruturados."""
    blocks = []
    lines = text.replace("\r\n", "\n").split("\n")
    i = 0
    para_buffer = []

    def flush_para():
        if para_buffer:
            joined = " ".join(l.strip() for l in para_buffer).strip()
            if joined:
                blocks.append(("p", joined))
            para_buffer.clear()

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Tabela: linha com "|" seguida imediatamente por uma linha separadora.
        if "|" in stripped and i + 1 < len(lines) and _is_separator_row(lines[i + 1]):
            flush_para()
            header = _split_row(lines[i])
            align = _parse_align(lines[i + 1])
            j = i + 2
            rows = []
            while j < len(lines) and lines[j].strip() and "|" in lines[j]:
                rows.append(_split_row(lines[j]))
                j += 1
            blocks.append(("table", {"header": header, "rows": rows, "align": align}))
            i = j
            continue

        if stripped == "":
            flush_para()
        elif stripped in ("---", "***", "___"):
            flush_para()
            blocks.append(("pagebreak", ""))
        elif stripped.startswith("### "):
            flush_para()
            blocks.append(("h3", stripped[4:].strip()))
        elif stripped.startswith("## "):
            flush_para()
            blocks.append(("h2", stripped[3:].strip()))
        elif stripped.startswith("# "):
            flush_para()
            blocks.append(("h1", stripped[2:].strip()))
        elif re.match(r"^[-*]\s+", stripped):
            flush_para()
            blocks.append(("ul", re.sub(r"^[-*]\s+", "", stripped)))
        elif re.match(r"^\d+[.)]\s+", stripped):
            flush_para()
            blocks.append(("ol", re.sub(r"^\d+[.)]\s+", "", stripped)))
        else:
            para_buffer.append(line)
        i += 1

    flush_para()
    return blocks


def split_bold(text):
    """Divide texto em segmentos (texto, is_bold) tratando **negrito**."""
    parts = []
    for idx, seg in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if seg == "":
            continue
        parts.append((seg, idx % 2 == 1))
    if not parts:
        parts = [(text, False)]
    return parts


# Cores da marca (ajuste aqui para casar com sua identidade visual).
BRAND_HEADER_HEX = "#2F5496"   # fundo do cabeçalho da tabela
BRAND_ZEBRA_HEX = "#F2F2F2"    # fundo das linhas alternadas
BRAND_GRID_HEX = "#BFBFBF"     # cor das linhas de grade


# ------------------------------- Geração DOCX --------------------------------

def _add_docx_table(doc, tdata):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import RGBColor

    header = tdata["header"]
    rows = tdata["rows"]
    align = tdata.get("align") or []
    ncols = max(len(header), max((len(r) for r in rows), default=0))

    align_map = {
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
    }

    def col_align(idx):
        a = align[idx] if idx < len(align) else ""
        return align_map.get(a)

    table = doc.add_table(rows=1, cols=ncols)
    # Estilo com grade; cai para 'Table Grid' se o tema não tiver o realçado.
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        try:
            table.style = "Table Grid"
        except KeyError:
            pass

    hdr_cells = table.rows[0].cells
    for idx in range(ncols):
        text = header[idx] if idx < len(header) else ""
        para = hdr_cells[idx].paragraphs[0]
        a = col_align(idx)
        if a is not None:
            para.alignment = a
        for seg, _ in split_bold(text):
            run = para.add_run(seg)
            run.bold = True  # cabeçalho sempre em negrito

    for row in rows:
        cells = table.add_row().cells
        for idx in range(ncols):
            val = row[idx] if idx < len(row) else ""
            para = cells[idx].paragraphs[0]
            a = col_align(idx)
            if a is not None:
                para.alignment = a
            for seg, is_bold in split_bold(val):
                run = para.add_run(seg)
                run.bold = is_bold

    doc.add_paragraph()  # respiro depois da tabela


def build_docx(blocks, out_path):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError:
        print("ERRO: falta 'python-docx'. Rode: pip install python-docx", file=sys.stderr)
        raise

    doc = Document()

    # Estilo base legível
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    def add_runs(paragraph, text):
        for seg, is_bold in split_bold(text):
            run = paragraph.add_run(seg)
            run.bold = is_bold

    for kind, content in blocks:
        if kind == "h1":
            p = doc.add_heading(level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_runs(p, content)
        elif kind == "h2":
            add_runs(doc.add_heading(level=1), content)
        elif kind == "h3":
            add_runs(doc.add_heading(level=2), content)
        elif kind == "ul":
            add_runs(doc.add_paragraph(style="List Bullet"), content)
        elif kind == "ol":
            add_runs(doc.add_paragraph(style="List Number"), content)
        elif kind == "table":
            _add_docx_table(doc, content)
        elif kind == "pagebreak":
            doc.add_page_break()
        elif kind == "p":
            add_runs(doc.add_paragraph(), content)

    doc.save(out_path)
    return out_path


# -------------------------------- Geração PDF --------------------------------

def _escape_rl(text):
    """Escapa para o mini-HTML do reportlab e converte **negrito**."""
    out = []
    for seg, is_bold in split_bold(text):
        seg = seg.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        out.append(f"<b>{seg}</b>" if is_bold else seg)
    return "".join(out)


def _register_unicode_font():
    """Registra uma fonte TrueType Unicode do sistema (acentos corretos).

    Retorna (regular_name, bold_name). Se nada for encontrado, cai para as
    fontes padrão Helvetica (que ainda cobrem o português, mas via WinAnsi)."""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # Pares (regular, bold) de caminhos candidatos, por sistema.
    candidates = [
        ("C:/Windows/Fonts/calibri.ttf", "C:/Windows/Fonts/calibrib.ttf"),
        ("C:/Windows/Fonts/arial.ttf", "C:/Windows/Fonts/arialbd.ttf"),
        ("C:/Windows/Fonts/segoeui.ttf", "C:/Windows/Fonts/segoeuib.ttf"),
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
         "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
        ("/Library/Fonts/Arial.ttf", "/Library/Fonts/Arial Bold.ttf"),
        ("/System/Library/Fonts/Supplemental/Arial.ttf",
         "/System/Library/Fonts/Supplemental/Arial Bold.ttf"),
    ]
    for reg, bold in candidates:
        if os.path.exists(reg):
            try:
                pdfmetrics.registerFont(TTFont("DocFont", reg))
                bold_name = "DocFont"
                if os.path.exists(bold):
                    pdfmetrics.registerFont(TTFont("DocFont-Bold", bold))
                    bold_name = "DocFont-Bold"
                    from reportlab.lib.fonts import addMapping
                    addMapping("DocFont", 0, 0, "DocFont")
                    addMapping("DocFont", 1, 0, "DocFont-Bold")
                return "DocFont", bold_name
            except Exception:
                continue
    return "Helvetica", "Helvetica-Bold"


def _make_pdf_table(tdata, font, font_bold, body_style, avail_width):
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import Paragraph, Table, TableStyle

    header = tdata["header"]
    rows = tdata["rows"]
    align = tdata.get("align") or []
    ncols = max(len(header), max((len(r) for r in rows), default=0))
    if ncols == 0:
        return None

    cell = ParagraphStyle("Cell", parent=body_style, fontName=font,
                          fontSize=10, leading=13, spaceAfter=0, spaceBefore=0)
    head = ParagraphStyle("CellHead", parent=cell, fontName=font_bold,
                          textColor=colors.white)

    # Células com Paragraph seguem o alinhamento DO PARÁGRAFO (não o do
    # TableStyle), então criamos uma variação de estilo por coluna.
    ta_map = {"right": TA_RIGHT, "center": TA_CENTER, "left": TA_LEFT}

    def styled(base, ci):
        a = align[ci] if ci < len(align) else ""
        if a in ta_map:
            return ParagraphStyle(f"{base.name}{ci}", parent=base,
                                  alignment=ta_map[a])
        return base

    def norm(row):
        return [row[i] if i < len(row) else "" for i in range(ncols)]

    data = [[Paragraph(_escape_rl(c), styled(head, ci))
             for ci, c in enumerate(norm(header))]]
    for row in rows:
        data.append([Paragraph(_escape_rl(c), styled(cell, ci))
                     for ci, c in enumerate(norm(row))])

    # Larguras proporcionais ao maior conteúdo de cada coluna, somando a
    # largura disponível — evita estourar a margem da página.
    weights = []
    for ci in range(ncols):
        longest = len(header[ci]) if ci < len(header) else 1
        for row in rows:
            if ci < len(row):
                longest = max(longest, len(row[ci]))
        weights.append(max(longest, 1))
    total = float(sum(weights)) or 1.0
    col_widths = [avail_width * w / total for w in weights]

    tbl = Table(data, colWidths=col_widths, hAlign="LEFT", repeatRows=1)
    style = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(BRAND_HEADER_HEX)),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor(BRAND_GRID_HEX)),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
         [colors.white, colors.HexColor(BRAND_ZEBRA_HEX)]),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]
    for ci in range(ncols):
        a = align[ci] if ci < len(align) else ""
        if a:
            style.append(("ALIGN", (ci, 0), (ci, -1), a.upper()))
    tbl.setStyle(TableStyle(style))
    return tbl


def build_pdf(blocks, out_path):
    try:
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import (ListFlowable, ListItem, PageBreak,
                                        Paragraph, SimpleDocTemplate, Spacer)
    except ImportError:
        print("ERRO: falta 'reportlab'. Rode: pip install reportlab", file=sys.stderr)
        raise

    font, font_bold = _register_unicode_font()

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("DocTitle", parent=styles["Title"],
                              alignment=TA_CENTER, fontSize=20, spaceAfter=18,
                              fontName=font_bold))
    body = ParagraphStyle("Body", parent=styles["BodyText"],
                          fontSize=11, leading=16, spaceAfter=8, fontName=font)
    h2 = ParagraphStyle("H2", parent=styles["Heading1"], fontSize=15,
                        spaceBefore=12, spaceAfter=6, fontName=font_bold)
    h3 = ParagraphStyle("H3", parent=styles["Heading2"], fontSize=12.5,
                        spaceBefore=8, spaceAfter=4, fontName=font_bold)

    left_margin = right_margin = 2.2 * cm
    doc = SimpleDocTemplate(out_path, pagesize=A4,
                            leftMargin=left_margin, rightMargin=right_margin,
                            topMargin=2 * cm, bottomMargin=2 * cm)
    avail_width = A4[0] - left_margin - right_margin
    story = []
    pending = []  # itens de lista acumulados

    def flush_list(kind):
        if pending:
            bt = "1" if kind == "ol" else "bullet"
            story.append(ListFlowable(
                [ListItem(Paragraph(_escape_rl(t), body)) for t in pending],
                bulletType=bt, leftIndent=18))
            pending.clear()

    prev_list = None
    for kind, content in blocks:
        if kind in ("ul", "ol"):
            if prev_list and prev_list != kind:
                flush_list(prev_list)
            pending.append(content)
            prev_list = kind
            continue
        if prev_list:
            flush_list(prev_list)
            prev_list = None

        if kind == "h1":
            story.append(Paragraph(_escape_rl(content), styles["DocTitle"]))
        elif kind == "h2":
            story.append(Paragraph(_escape_rl(content), h2))
        elif kind == "h3":
            story.append(Paragraph(_escape_rl(content), h3))
        elif kind == "p":
            story.append(Paragraph(_escape_rl(content), body))
        elif kind == "table":
            tbl = _make_pdf_table(content, font, font_bold, body, avail_width)
            if tbl is not None:
                story.append(tbl)
                story.append(Spacer(1, 10))
        elif kind == "pagebreak":
            story.append(PageBreak())

    if prev_list:
        flush_list(prev_list)

    if not story:
        story.append(Spacer(1, 1))
    doc.build(story)
    return out_path


# ----------------------------------- Main ------------------------------------

def main():
    ap = argparse.ArgumentParser(description="Gera .docx e/ou .pdf a partir de Markdown.")
    ap.add_argument("input", help="Arquivo .md de entrada")
    ap.add_argument("--out", help="Nome base de saída (sem extensão)")
    ap.add_argument("--outdir", default=".", help="Pasta de saída")
    ap.add_argument("--docx", action="store_true", help="Gerar .docx")
    ap.add_argument("--pdf", action="store_true", help="Gerar .pdf")
    args = ap.parse_args()

    with open(args.input, "r", encoding="utf-8") as f:
        text = f.read()

    blocks = parse_markdown(text)
    if not blocks:
        print("AVISO: nenhum conteúdo encontrado no markdown.", file=sys.stderr)

    base = args.out or os.path.splitext(os.path.basename(args.input))[0]
    os.makedirs(args.outdir, exist_ok=True)

    want_docx, want_pdf = args.docx, args.pdf
    if not want_docx and not want_pdf:
        want_docx = want_pdf = True

    created = []
    if want_docx:
        p = os.path.join(args.outdir, base + ".docx")
        build_docx(blocks, p)
        created.append(p)
    if want_pdf:
        p = os.path.join(args.outdir, base + ".pdf")
        build_pdf(blocks, p)
        created.append(p)

    for p in created:
        print("Gerado:", os.path.abspath(p))


if __name__ == "__main__":
    main()
